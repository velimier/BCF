from __future__ import annotations
import json, os, random, statistics, threading
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor, as_completed
from datetime import datetime
from dataclasses import asdict
from itertools import combinations
from typing import Dict, List, Tuple, Optional

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QTabWidget, QVBoxLayout, QGridLayout, QLabel,
    QLineEdit, QPushButton, QFileDialog, QSpinBox, QDoubleSpinBox, QCheckBox, QMessageBox,
    QHBoxLayout, QComboBox, QGroupBox, QFrame, QScrollArea
)
from PySide6.QtGui import QAction
from PySide6.QtCore import Signal, QObject, Qt

from matplotlib.backends.backend_qtagg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.backends.backend_qtagg import NavigationToolbar2QT as NavigationToolbar
from matplotlib.figure import Figure
from matplotlib.ticker import ScalarFormatter, StrMethodFormatter

from ..engine.models import RockMass, JointSet, SpacingDist, CaveFace, Defaults, SecondaryRun, PrimaryBlock
from ..engine.primary import generate_primary_blocks
from ..engine.secondary import run_secondary, average_scatter_deg_from_jointsets
from ..engine.hangup import orepass_hangups, kear_hangups
from ..engine.io_formats import distributions_from_blocks, write_prm, write_sec


def randomize_value(
    value: float,
    variation_pct: float,
    minimum: float | None = None,
    maximum: float | None = None,
    rng: random.Random | None = None,
) -> float:
    rng = rng or random
    if variation_pct <= 0:
        new_val = value
    else:
        delta = variation_pct / 100.0
        new_val = value * (1.0 + rng.uniform(-delta, delta))
    if minimum is not None:
        new_val = max(minimum, new_val)
    if maximum is not None:
        new_val = min(maximum, new_val)
    return new_val


def randomize_joint_set(js: JointSet, variation_pct: float, rng: random.Random | None = None) -> JointSet:
    rng = rng or random
    spacing_vals = [
        randomize_value(js.spacing.min, variation_pct, 0.01, rng=rng),
        randomize_value(js.spacing.mean, variation_pct, 0.02, rng=rng),
        randomize_value(js.spacing.max_or_90pct, variation_pct, 0.05, rng=rng),
    ]
    spacing_vals.sort()
    spacing = SpacingDist(
        js.spacing.type,
        spacing_vals[0],
        max(spacing_vals[0], spacing_vals[1]),
        max(spacing_vals[1], spacing_vals[2])
    )
    return JointSet(
        name=js.name,
        mean_dip=randomize_value(js.mean_dip, variation_pct, 0.0, 90.0, rng=rng),
        dip_range=randomize_value(js.dip_range, variation_pct, 0.0, 90.0, rng=rng),
        mean_dip_dir=randomize_value(js.mean_dip_dir, variation_pct, 0.0, 360.0, rng=rng),
        dip_dir_range=randomize_value(js.dip_dir_range, variation_pct, 0.0, 180.0, rng=rng),
        spacing=spacing,
        JC=int(round(randomize_value(js.JC, variation_pct, 0, 40, rng=rng)))
    )


MC_COLOR_OPTIONS: List[Tuple[str, str]] = [
    ("Blue", "#1f77b4"),
    ("Orange", "#ff7f0e"),
    ("Green", "#2ca02c"),
    ("Red", "#d62728"),
    ("Purple", "#9467bd"),
    ("Brown", "#8c564b"),
    ("Pink", "#e377c2"),
    ("Gray", "#7f7f7f"),
    ("Olive", "#bcbd22"),
    ("Cyan", "#17becf"),
    ("Black", "#000000"),
]

LINE_STYLE_OPTIONS: List[Tuple[str, str]] = [
    ("Solid", "-"),
    ("Dashed", "--"),
    ("Dotted", ":"),
    ("Dash-dot", "-.")
]


def _monte_carlo_worker(
    seed: int,
    nblocks: int,
    variation: float,
    base_inputs: Dict[str, object],
    combination: Optional[Tuple[int, int, int]] = None,
    combo_label: Optional[str] = None,
) -> Tuple[str, dict, dict]:
    rng = random.Random(seed)

    rock_dict = base_inputs.get("rock", {})
    rock = RockMass(
        rock_type=rock_dict.get("rock_type", "Unknown"),
        MRMR=randomize_value(rock_dict.get("MRMR", 65.0), variation, 0.0, 100.0, rng=rng),
        IRS=randomize_value(rock_dict.get("IRS", 120.0), variation, 1.0, 500.0, rng=rng),
        IBS=rock_dict.get("IBS"),
        mi=randomize_value(rock_dict.get("mi", 17.0), variation, 1.0, 50.0, rng=rng),
        frac_freq=randomize_value(rock_dict.get("frac_freq", 0.0), variation, 0.0, 20.0, rng=rng),
        frac_condition=int(round(randomize_value(rock_dict.get("frac_condition", 20), variation, 0, 40, rng=rng))),
        density=randomize_value(rock_dict.get("density", 3200.0), variation, 1500.0, 4500.0, rng=rng),
    )

    randomized_sets: List[JointSet] = []
    for js_dict in base_inputs.get("joint_sets", []):
        spacing_dict = js_dict.get("spacing", {})
        spacing = SpacingDist(
            spacing_dict.get("type", "trunc_exp"),
            spacing_dict.get("min", 0.3),
            spacing_dict.get("mean", 1.0),
            spacing_dict.get("max_or_90pct", 3.0),
            spacing_dict.get("max_obs"),
        )
        js = JointSet(
            name=js_dict.get("name", "Set"),
            mean_dip=js_dict.get("mean_dip", 45.0),
            dip_range=js_dict.get("dip_range", 10.0),
            mean_dip_dir=js_dict.get("mean_dip_dir", 0.0),
            dip_dir_range=js_dict.get("dip_dir_range", 20.0),
            spacing=spacing,
            JC=js_dict.get("JC", 20),
        )
        randomized_sets.append(randomize_joint_set(js, variation, rng=rng))

    if combination is not None:
        joint_sets = [randomized_sets[i] for i in combination if 0 <= i < len(randomized_sets)]
    else:
        joint_sets = list(randomized_sets)

    if len(joint_sets) < 3:
        raise ValueError("Monte Carlo worker requires at least three joint sets.")

    cave_dict = base_inputs.get("cave", {})
    cave = CaveFace(
        dip=cave_dict.get("dip", 45.0),
        dip_dir=cave_dict.get("dip_dir", 0.0),
        stress_dip=cave_dict.get("stress_dip", 5.0),
        stress_strike=cave_dict.get("stress_strike", 5.0),
        stress_normal=cave_dict.get("stress_normal", 0.0),
        allow_stress_fractures=cave_dict.get("allow_stress_fractures", True),
        spalling_pct=randomize_value(cave_dict.get("spalling_pct", 0.0), variation, 0.0, 100.0, rng=rng),
    )

    defaults_dict = base_inputs.get("defaults", {})
    defaults = Defaults(
        LHD_cutoff_m3=randomize_value(defaults_dict.get("LHD_cutoff_m3", 2.0), variation, 0.1, 50.0, rng=rng),
        seed=rng.randint(0, 10**9),
        arching_pct=randomize_value(defaults_dict.get("arching_pct", 0.12), variation, 0.0, 1.0, rng=rng),
        arch_stress_conc=randomize_value(defaults_dict.get("arch_stress_conc", 25.0), variation, 1.0, 100.0, rng=rng),
        stress_frac_trace_min=defaults_dict.get("stress_frac_trace_min", 10.0),
        stress_frac_trace_mean=defaults_dict.get("stress_frac_trace_mean", 20.0),
        stress_frac_trace_max=defaults_dict.get("stress_frac_trace_max", 30.0),
        tension_factor=defaults_dict.get("tension_factor", 0.0),
    )

    secondary_dict = base_inputs.get("secondary", {})
    secondary_params = SecondaryRun(
        draw_height=randomize_value(secondary_dict.get("draw_height", 150.0), variation, 1.0, 2000.0, rng=rng),
        max_caving_height=randomize_value(secondary_dict.get("max_caving_height", 300.0), variation, 1.0, 5000.0, rng=rng),
        swell_factor=randomize_value(secondary_dict.get("swell_factor", 1.2), variation, 1.0, 3.0, rng=rng),
        active_draw_width=randomize_value(secondary_dict.get("active_draw_width", 45.0), variation, 1.0, 200.0, rng=rng),
        added_fines_pct=randomize_value(secondary_dict.get("added_fines_pct", 0.0), variation, 0.0, 80.0, rng=rng),
        rate_cm_day=randomize_value(secondary_dict.get("rate_cm_day", 20.0), variation, 0.0, 100.0, rng=rng),
        drawbell_upper_width=randomize_value(secondary_dict.get("drawbell_upper_width", 8.0), variation, 1.0, 50.0, rng=rng),
        drawbell_lower_width=randomize_value(secondary_dict.get("drawbell_lower_width", 6.0), variation, 1.0, 50.0, rng=rng),
    )

    blocks = generate_primary_blocks(nblocks, rock, joint_sets, cave, defaults, seed=defaults.seed)
    primary_stats = distributions_from_blocks(blocks)
    mu = average_scatter_deg_from_jointsets(joint_sets)
    primary_fines_ratio = cave.spalling_pct / 100.0
    sec_blocks, _ = run_secondary(
        blocks,
        rock,
        secondary_params,
        defaults,
        mu_scatter_deg=mu,
        primary_fines_ratio=primary_fines_ratio,
    )
    secondary_stats = distributions_from_blocks(sec_blocks)
    label = combo_label or "+".join(js.name for js in joint_sets[:3])
    return label, primary_stats, secondary_stats


class PlotWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.fig = Figure(figsize=(6, 4))
        self.canvas = FigureCanvas(self.fig)
        self.toolbar = NavigationToolbar(self.canvas, self)
        lay = QVBoxLayout()
        lay.setContentsMargins(8, 8, 8, 8)
        lay.setSpacing(8)
        lay.addWidget(self.toolbar)
        lay.addWidget(self.canvas)
        self.setLayout(lay)
        self.toolbar.setStyleSheet("QToolBar { border: 0; background: transparent; }")
        self.canvas.setMinimumHeight(340)
        self.canvas.setStyleSheet(
            "background-color: #ffffff; border: 1px solid #d4d9e2; border-radius: 8px;"
        )
        self.fig.set_facecolor("#f3f6fb")
        self.has_data = False

    def plot_distributions(self, prim_stats: dict, sec_stats: dict | None = None, title: str = ""):
        self.fig.clear()
        ax = self.fig.add_subplot(111)
        xs = [lo for lo,hi in prim_stats["bins"]]
        ax.plot(xs, prim_stats["cum_mass"], label="Primary")
        if sec_stats is not None:
            ax.plot(xs, sec_stats["cum_mass"], label="Secondary")
        ax.set_xscale("log")
        ax.set_xlabel("Block volume (m³)")
        ax.set_ylabel("Cumulative mass (%)")
        ax.set_title(title)
        ax.grid(True, which="both", linestyle=":", color="#cfd6e6", linewidth=0.6)
        self._finalize_axes(ax, show_legend=True)
        self.canvas.draw_idle()
        self.has_data = True

    def plot_lines(
        self,
        xs: List[float] | List[List[float]],
        ys_list: List[List[float]],
        labels: List[str] | None = None,
        title: str = "",
        xlabel: str = "",
        ylabel: str = "",
        logx: bool = False,
        styles: List[Dict[str, object]] | None = None,
        x_formatter=None,
        y_formatter=None,
    ):
        self.fig.clear()
        ax = self.fig.add_subplot(111)
        if ys_list:
            if isinstance(xs, list) and xs and isinstance(xs[0], (list, tuple)):
                xs_list = list(xs)
            else:
                xs_list = [xs] * len(ys_list)
        else:
            xs_list = []
        style_lookup: Dict[str, Dict[str, object]] = {}
        style_list: List[Dict[str, object]] = []
        if isinstance(styles, dict):
            style_lookup = styles
        elif isinstance(styles, list):
            style_list = styles
        for i, ys in enumerate(ys_list):
            label = labels[i] if labels and i < len(labels) else None
            cur_xs = xs_list[i] if i < len(xs_list) else xs_list[0] if xs_list else []
            if label is not None and style_lookup:
                style = style_lookup.get(label, {})
            elif style_list and i < len(style_list):
                style = style_list[i]
            else:
                style = {}
            display_label = style.get("label") if isinstance(style, dict) and style.get("label") else label
            color = style.get("color") if isinstance(style, dict) else None
            linestyle = style.get("linestyle") if isinstance(style, dict) and style.get("linestyle") else "-"
            linewidth = style.get("linewidth") if isinstance(style, dict) and style.get("linewidth") else 1.5
            ax.plot(cur_xs, ys, label=display_label, color=color, linestyle=linestyle, linewidth=linewidth)
        if logx:
            ax.set_xscale("log")
        ax.set_xlabel(xlabel)
        ax.set_ylabel(ylabel)
        ax.set_title(title)
        ax.grid(True, which="both", linestyle=":", color="#cfd6e6", linewidth=0.6)
        show_legend = bool(ax.get_legend_handles_labels()[1])
        self._finalize_axes(ax, show_legend=show_legend)
        if x_formatter is not None:
            ax.xaxis.set_major_formatter(x_formatter)
        if y_formatter is not None:
            ax.yaxis.set_major_formatter(y_formatter)
        self.canvas.draw_idle()
        self.has_data = True

    def _finalize_axes(self, ax, *, show_legend: bool):
        for spine in ("top", "right"):
            if spine in ax.spines:
                ax.spines[spine].set_visible(False)
        for spine in ("bottom", "left"):
            if spine in ax.spines:
                ax.spines[spine].set_color("#c4ccdc")
                ax.spines[spine].set_linewidth(0.8)
        ax.tick_params(axis="both", colors="#2e3650", labelsize=9)
        ax.set_facecolor("#ffffff")
        legend = None
        if show_legend:
            handles, legend_labels = ax.get_legend_handles_labels()
            if legend_labels:
                legend = ax.legend(
                    handles,
                    legend_labels,
                    loc="center left",
                    bbox_to_anchor=(1.02, 0.5),
                    frameon=False,
                    fontsize=9,
                    labelspacing=0.6,
                )
        ax.figure.set_facecolor("#f3f6fb")
        if show_legend and legend is not None:
            ax.figure.subplots_adjust(left=0.12, right=0.74, top=0.88, bottom=0.16)
        else:
            ax.figure.subplots_adjust(left=0.12, right=0.96, top=0.88, bottom=0.16)

    def save_dialog(self, parent: QWidget, suggested_name: str = "chart.png"):
        if not self.has_data:
            QMessageBox.warning(parent, "No chart", "There is no chart to save yet.")
            return
        path, _ = QFileDialog.getSaveFileName(parent, "Save chart", suggested_name,
                                             "PNG image (*.png);;PDF document (*.pdf);;SVG image (*.svg)")
        if path:
            try:
                self.fig.savefig(path, dpi=300, bbox_inches="tight")
            except Exception as exc:  # pragma: no cover - UI feedback
                QMessageBox.critical(parent, "Save failed", f"Could not save chart:\n{exc}")
            else:
                QMessageBox.information(parent, "Chart saved", f"Chart saved to:\n{path}")

class AppSignals(QObject):
    log = Signal(str)
    done_primary = Signal(list, float, str)
    done_secondary = Signal(list, float, str)
    done_hangup = Signal(dict)
    done_monte_carlo = Signal(dict)

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("BCF-Style Fragmentation (Python, PySide6)")
        self.resize(1100, 760)
        self.sig = AppSignals()

        self.tabs = QTabWidget()
        self.tabs.setDocumentMode(True)
        self.setCentralWidget(self.tabs)
        self._apply_base_styles()

        file_menu = self.menuBar().addMenu("&File")
        self.action_save_settings = QAction("Save settings…", self)
        self.action_load_settings = QAction("Load settings…", self)
        file_menu.addAction(self.action_save_settings)
        file_menu.addAction(self.action_load_settings)

        self.rock = RockMass()
        self.joint_sets: List[JointSet] = [
            JointSet("Set1", 60, 20, 90, 20, SpacingDist("trunc_exp",0.3,1.0,3.0), 20),
            JointSet("Set2", 45, 20, 0,  20, SpacingDist("trunc_exp",0.3,1.5,4.0), 20),
            JointSet("Set3", 30, 20, 180,20, SpacingDist("trunc_exp",0.4,2.0,5.0), 20),
        ]
        self.cave = CaveFace()
        self.defaults = Defaults()
        self.secondary = SecondaryRun()

        self.primary_blocks: List[PrimaryBlock] = []
        self.primary_fines_ratio = 0.0
        self.secondary_blocks = []
        self._last_monte_carlo_result: dict | None = None
        self._last_monte_carlo_inputs: dict | None = None
        self._last_monte_carlo_settings: dict | None = None
        self._last_primary_stats: dict | None = None
        self._last_secondary_stats: dict | None = None

        self._series_style_widgets: Dict[str, Dict[str, QWidget]] = {}
        self._series_color_cursor = 0
        self._chart_title_widgets: Dict[str, QLineEdit] = {}
        self._axis_format_combos: Dict[str, QComboBox] = {}
        self._series_random_defaults: Dict[str, Tuple[int, int, float]] = {}
        self._combo_style_cache: Dict[str, dict] = {}

        self._build_tabs()
        self._connect_signals()
        self._ensure_series_style_controls(["Primary", "Secondary"])
        self._update_combination_controls()

    def _apply_base_styles(self):
        self.setStyleSheet(
            """
            QMainWindow {
                background-color: #e9edf4;
            }
            QWidget {
                font-size: 11pt;
            }
            QTabWidget::pane {
                border: 1px solid #c8cfdd;
                border-radius: 6px;
                background: #f9fbfe;
            }
            QTabBar::tab {
                padding: 8px 18px;
                background: #dfe5f1;
                border-top-left-radius: 6px;
                border-top-right-radius: 6px;
            }
            QTabBar::tab:selected {
                background: #ffffff;
                color: #1a2132;
            }
            QScrollArea {
                background: transparent;
            }
            QFrame[frameShape="StyledPanel"] {
                background-color: #ffffff;
                border: 1px solid #d4d9e2;
                border-radius: 8px;
            }
            QPushButton {
                background-color: #2f6fed;
                color: #ffffff;
                border-radius: 6px;
                padding: 6px 14px;
            }
            QPushButton:hover:!disabled {
                background-color: #255ad0;
            }
            QPushButton:disabled {
                background-color: #b8c2d6;
                color: #f3f4f8;
            }
            QLabel {
                color: #1f2537;
            }
            QGroupBox {
                border: 1px solid #d4d9e2;
                border-radius: 8px;
                margin-top: 12px;
                background: #ffffff;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 16px;
                padding: 0 4px;
                background: transparent;
            }
            QToolButton {
                background: transparent;
                border: none;
            }
            QLineEdit, QSpinBox, QDoubleSpinBox, QComboBox {
                background: #ffffff;
                border: 1px solid #c8cfdd;
                border-radius: 4px;
                padding: 2px 6px;
            }
            QLineEdit:focus, QSpinBox:focus, QDoubleSpinBox:focus, QComboBox:focus {
                border: 1px solid #7b95f2;
            }
            """
        )

    def _set_uniform_input_width(self, widget):
        widget.setMinimumWidth(140)
        widget.setMaximumWidth(220)

    def _joint_widget_display_name(self, index: int) -> str:
        if 0 <= index < len(self.joint_widgets):
            name_widget = self.joint_widgets[index].get("name")
            if isinstance(name_widget, QLineEdit):
                text = name_widget.text().strip()
                if text:
                    return text
        return f"Set{index + 1}"

    def _refresh_joint_remove_buttons(self):
        allow_remove = len(self.joint_widgets) > 3
        for entry in self.joint_widgets:
            btn = entry.get("remove")
            if isinstance(btn, QPushButton):
                btn.setEnabled(allow_remove)

    def _add_joint_set_widget(self, joint_set: JointSet | None = None, *, suppress_update: bool = False):
        js = joint_set or JointSet(
            name=f"Set{len(self.joint_widgets) + 1}",
            mean_dip=45.0,
            dip_range=10.0,
            mean_dip_dir=0.0,
            dip_dir_range=20.0,
            spacing=SpacingDist("trunc_exp", 0.3, 1.0, 3.0),
            JC=20,
        )
        frame = QFrame()
        frame.setFrameShape(QFrame.StyledPanel)
        layout = QGridLayout(frame)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setHorizontalSpacing(10)
        layout.setVerticalSpacing(6)

        name_edit = QLineEdit(js.name)
        self._set_uniform_input_width(name_edit)
        remove_btn = QPushButton("Remove")
        remove_btn.setAutoDefault(False)
        remove_btn.setDefault(False)

        dip = QDoubleSpinBox(); dip.setRange(0, 90); dip.setValue(js.mean_dip); self._set_uniform_input_width(dip)
        dipr = QDoubleSpinBox(); dipr.setRange(0, 90); dipr.setValue(js.dip_range); self._set_uniform_input_width(dipr)
        dd = QDoubleSpinBox(); dd.setRange(0, 360); dd.setValue(js.mean_dip_dir); self._set_uniform_input_width(dd)
        ddr = QDoubleSpinBox(); ddr.setRange(0, 180); ddr.setValue(js.dip_dir_range); self._set_uniform_input_width(ddr)
        jc = QSpinBox(); jc.setRange(0, 40); jc.setValue(js.JC); self._set_uniform_input_width(jc)

        spacing_min = getattr(js.spacing, "min", 0.3)
        spacing_mean = getattr(js.spacing, "mean", spacing_min)
        spacing_max = getattr(js.spacing, "max_or_90pct", max(spacing_mean, spacing_min))

        s_min = QDoubleSpinBox(); s_min.setRange(0.01, 50); s_min.setDecimals(2); s_min.setValue(spacing_min); self._set_uniform_input_width(s_min)
        s_mean = QDoubleSpinBox(); s_mean.setRange(0.02, 50); s_mean.setDecimals(2); s_mean.setValue(spacing_mean); self._set_uniform_input_width(s_mean)
        s_max = QDoubleSpinBox(); s_max.setRange(0.03, 200); s_max.setDecimals(2); s_max.setValue(spacing_max); self._set_uniform_input_width(s_max)

        layout.addWidget(QLabel("Name"), 0, 0)
        layout.addWidget(name_edit, 0, 1)
        layout.addWidget(remove_btn, 0, 2)
        layout.addWidget(QLabel("Dip / Range"), 1, 0)
        layout.addWidget(dip, 1, 1)
        layout.addWidget(dipr, 1, 2)
        layout.addWidget(QLabel("Dip dir / Range"), 2, 0)
        layout.addWidget(dd, 2, 1)
        layout.addWidget(ddr, 2, 2)
        layout.addWidget(QLabel("JC (0–40)"), 3, 0)
        layout.addWidget(jc, 3, 1)
        layout.addWidget(QLabel("Spacing min / mean / max"), 4, 0)
        layout.addWidget(s_min, 4, 1)
        layout.addWidget(s_mean, 4, 2)
        layout.addWidget(s_max, 4, 3)
        layout.setColumnStretch(1, 1)
        layout.setColumnStretch(2, 1)
        layout.setColumnStretch(3, 1)

        entry: Dict[str, QWidget] = {}
        entry.update({
            "frame": frame,
            "name": name_edit,
            "dip": dip,
            "dip_range": dipr,
            "dip_dir": dd,
            "dip_dir_range": ddr,
            "jc": jc,
            "spacing_min": s_min,
            "spacing_mean": s_mean,
            "spacing_max": s_max,
            "remove": remove_btn,
        })

        def handle_remove():
            self._remove_joint_set_widget(entry)

        def handle_name_change():
            frame.setProperty("joint_set_name", name_edit.text())
            self._update_combination_controls()

        remove_btn.clicked.connect(handle_remove)
        name_edit.editingFinished.connect(self._update_combination_controls)
        name_edit.textChanged.connect(handle_name_change)

        self.joint_widgets.append(entry)
        if hasattr(self, "joint_sets_container"):
            self.joint_sets_container.addWidget(frame)
        if not suppress_update and hasattr(self, "joint_sets_scroll"):
            scroll = getattr(self, "joint_sets_scroll", None)
            if isinstance(scroll, QScrollArea):
                bar = scroll.verticalScrollBar()
                if bar is not None:
                    bar.setValue(bar.maximum())
        self._refresh_joint_remove_buttons()
        if not suppress_update:
            self._update_combination_controls()

    def _remove_joint_set_widget(self, entry: Dict[str, QWidget]):
        if entry not in self.joint_widgets:
            return
        if len(self.joint_widgets) <= 3:
            QMessageBox.warning(self, "Cannot remove", "At least three joint sets are required for analysis.")
            return
        self.joint_widgets.remove(entry)
        frame = entry.get("frame")
        if isinstance(frame, QWidget):
            frame.setParent(None)
            frame.deleteLater()
        self._refresh_joint_remove_buttons()
        self._update_combination_controls()

    def _clear_joint_set_widgets(self):
        for entry in getattr(self, "joint_widgets", []):
            frame = entry.get("frame")
            if isinstance(frame, QWidget):
                frame.setParent(None)
                frame.deleteLater()
        self.joint_widgets = []

    def _rebuild_joint_set_widgets(self, joint_sets: List[JointSet]):
        self._clear_joint_set_widgets()
        for js in joint_sets:
            self._add_joint_set_widget(js, suppress_update=True)
        self._refresh_joint_remove_buttons()
        self._update_combination_controls()

    def _update_joint_combination_controls(self):
        combo = getattr(self, "combo_joint_combination", None)
        if not isinstance(combo, QComboBox):
            return
        current_data = combo.currentData()
        combo.blockSignals(True)
        combo.clear()
        count = len(self.joint_widgets)
        available = list(combinations(range(count), 3)) if count >= 3 else []
        if available:
            for idxs in available:
                label = " + ".join(self._joint_widget_display_name(i) for i in idxs)
                combo.addItem(label, idxs)
            if current_data in available:
                combo.setCurrentIndex(available.index(current_data))
            else:
                combo.setCurrentIndex(0)
            combo.setEnabled(True)
        else:
            combo.addItem("Add at least three joint sets", None)
            combo.setEnabled(False)
        combo.blockSignals(False)

    def _update_monte_carlo_combination_controls(self):
        combo = getattr(self, "combo_mc_combinations", None)
        if not isinstance(combo, QComboBox):
            return
        current_data = combo.currentData()
        combo.blockSignals(True)
        combo.clear()
        count = len(self.joint_widgets)
        available = list(combinations(range(count), 3)) if count >= 3 else []
        combo.addItem("Use geology selection", ("geology", None))
        for idxs in available:
            label = " + ".join(self._joint_widget_display_name(i) for i in idxs)
            combo.addItem(label, ("combo", idxs))
        if len(available) > 1:
            combo.addItem("All combinations", ("all", None))
        if current_data is not None:
            idx = combo.findData(current_data)
            if idx >= 0:
                combo.setCurrentIndex(idx)
        combo.blockSignals(False)

    def _update_combination_controls(self):
        self._update_joint_combination_controls()
        self._update_monte_carlo_combination_controls()

    def _selected_joint_combination_indexes(self) -> Optional[Tuple[int, int, int]]:
        combo = getattr(self, "combo_joint_combination", None)
        if not isinstance(combo, QComboBox):
            return None
        data = combo.currentData()
        if isinstance(data, tuple) and len(data) == 3:
            return data
        return None

    def _selected_joint_sets(self) -> List[JointSet]:
        idxs = self._selected_joint_combination_indexes()
        if not idxs:
            return []
        return [self.joint_sets[i] for i in idxs if 0 <= i < len(self.joint_sets)]

    def _combination_label(self, indexes: Tuple[int, int, int], names: Optional[List[str]] = None) -> str:
        if names is None:
            names = [js.name for js in self.joint_sets]
        parts = []
        for idx in indexes:
            if 0 <= idx < len(names):
                part = names[idx] or f"Set{idx + 1}"
            else:
                part = f"Set{idx + 1}"
            parts.append(part)
        return " + ".join(parts)

    def _all_joint_combinations(self) -> List[Tuple[int, int, int]]:
        return list(combinations(range(len(self.joint_sets)), 3)) if len(self.joint_sets) >= 3 else []

    def _compact_combo_label(self, label: str) -> str:
        parts = [part.strip() for part in label.split("+") if part.strip()]
        compact_parts: List[str] = []
        for part in parts:
            tokens = [tok for tok in part.replace("-", " ").split() if tok]
            if len(tokens) >= 2:
                abbrev = "".join(token[0].upper() for token in tokens[:3])
                if len(abbrev) >= 2:
                    compact_parts.append(abbrev)
                    continue
            if len(part) > 12:
                compact_parts.append(part[:11] + "…")
            else:
                compact_parts.append(part)
        return " / ".join(compact_parts) if compact_parts else label

    def _build_tabs(self):
        self.tabs.addTab(self._build_geology_tab(), "Geology")
        self.tabs.addTab(self._build_cave_tab(), "Cave")
        self.tabs.addTab(self._build_primary_tab(), "Primary run")
        self.tabs.addTab(self._build_secondary_tab(), "Secondary & hang-up")
        self.tabs.addTab(self._build_monte_carlo_tab(), "Monte Carlo")
        self.tabs.addTab(self._build_chart_settings_tab(), "Chart settings")
        self.tabs.addTab(self._build_defaults_tab(), "Defaults")

    def _build_geology_tab(self):
        w = QWidget(); g = QGridLayout(w)
        row = 0; g.addWidget(QLabel("<b>Rock mass</b>"), row,0,1,2); row+=1
        self.rock_type = QLineEdit(self.rock.rock_type); self._set_uniform_input_width(self.rock_type); g.addWidget(QLabel("Rock type"), row,0); g.addWidget(self.rock_type,row,1); row+=1
        self.mrmr = QDoubleSpinBox(); self.mrmr.setRange(0,100); self.mrmr.setValue(self.rock.MRMR); self._set_uniform_input_width(self.mrmr); g.addWidget(QLabel("MRMR"), row,0); g.addWidget(self.mrmr,row,1); row+=1
        self.irs = QDoubleSpinBox(); self.irs.setRange(1,500); self.irs.setValue(self.rock.IRS); self._set_uniform_input_width(self.irs); g.addWidget(QLabel("IRS (MPa)"), row,0); g.addWidget(self.irs,row,1); row+=1
        self.mi = QDoubleSpinBox(); self.mi.setRange(1,50); self.mi.setValue(self.rock.mi); self._set_uniform_input_width(self.mi); g.addWidget(QLabel("mi (Hoek–Brown)"), row,0); g.addWidget(self.mi,row,1); row+=1
        self.ff = QDoubleSpinBox(); self.ff.setRange(0,20); self.ff.setDecimals(2); self.ff.setValue(self.rock.frac_freq); self._set_uniform_input_width(self.ff); g.addWidget(QLabel("Fracture/veinlet freq (1/m)"), row,0); g.addWidget(self.ff,row,1); row+=1
        self.fc = QSpinBox(); self.fc.setRange(0,40); self.fc.setValue(self.rock.frac_condition); self._set_uniform_input_width(self.fc); g.addWidget(QLabel("Fracture/veinlet condition (0–40)"), row,0); g.addWidget(self.fc,row,1); row+=1
        self.density = QDoubleSpinBox(); self.density.setRange(1500,4500); self.density.setValue(self.rock.density); self._set_uniform_input_width(self.density); g.addWidget(QLabel("Density (kg/m³)"), row,0); g.addWidget(self.density,row,1); row+=1

        row += 1
        g.addWidget(QLabel("<b>Joint sets</b>"), row, 0, 1, 2)
        row += 1
        self.joint_widgets = []
        joint_sets_widget = QWidget()
        self.joint_sets_container = QVBoxLayout()
        self.joint_sets_container.setContentsMargins(0, 0, 0, 0)
        self.joint_sets_container.setSpacing(12)
        joint_sets_widget.setLayout(self.joint_sets_container)
        self.joint_sets_scroll = QScrollArea()
        self.joint_sets_scroll.setWidgetResizable(True)
        self.joint_sets_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.joint_sets_scroll.setFrameShape(QFrame.NoFrame)
        self.joint_sets_scroll.setWidget(joint_sets_widget)
        g.addWidget(self.joint_sets_scroll, row, 0, 1, 4)
        row += 1
        for js in self.joint_sets:
            self._add_joint_set_widget(js, suppress_update=True)

        btn_row = QHBoxLayout()
        self.btn_add_joint_set = QPushButton("Add joint set")
        btn_row.addWidget(self.btn_add_joint_set)
        btn_row.addStretch(1)
        g.addLayout(btn_row, row, 0, 1, 4)
        row += 1

        combo_row = QHBoxLayout()
        combo_row.addWidget(QLabel("Joint set combination (3 sets):"))
        self.combo_joint_combination = QComboBox()
        combo_row.addWidget(self.combo_joint_combination)
        combo_row.addStretch(1)
        g.addLayout(combo_row, row, 0, 1, 4)
        row += 1

        self._update_joint_combination_controls()

        w.setLayout(g)
        return w

    def _build_cave_tab(self):
        w = QWidget(); g = QGridLayout(w)
        row=0; g.addWidget(QLabel("<b>Cave face & stresses</b>"), row,0,1,2); row+=1
        self.cave_dip = QDoubleSpinBox(); self.cave_dip.setRange(0,90); self.cave_dip.setValue(self.cave.dip); self._set_uniform_input_width(self.cave_dip); g.addWidget(QLabel("Face dip (°)"), row,0); g.addWidget(self.cave_dip,row,1); row+=1
        self.cave_ddir = QDoubleSpinBox(); self.cave_ddir.setRange(0,360); self.cave_ddir.setValue(self.cave.dip_dir); self._set_uniform_input_width(self.cave_ddir); g.addWidget(QLabel("Face dip direction (°)"), row,0); g.addWidget(self.cave_ddir,row,1); row+=1
        self.st_dip = QDoubleSpinBox(); self.st_dip.setRange(0,100); self.st_dip.setValue(self.cave.stress_dip); self._set_uniform_input_width(self.st_dip); g.addWidget(QLabel("Dip stress (MPa)"), row,0); g.addWidget(self.st_dip,row,1); row+=1
        self.st_strike = QDoubleSpinBox(); self.st_strike.setRange(0,100); self.st_strike.setValue(self.cave.stress_strike); self._set_uniform_input_width(self.st_strike); g.addWidget(QLabel("Strike stress (MPa)"), row,0); g.addWidget(self.st_strike,row,1); row+=1
        self.st_norm = QDoubleSpinBox(); self.st_norm.setRange(0,100); self.st_norm.setValue(self.cave.stress_normal); self._set_uniform_input_width(self.st_norm); g.addWidget(QLabel("Normal stress (MPa)"), row,0); g.addWidget(self.st_norm,row,1); row+=1
        self.allow_sf = QCheckBox("Allow stress fractures"); self.allow_sf.setChecked(self.cave.allow_stress_fractures); g.addWidget(self.allow_sf,row,0,1,2); row+=1
        self.spalling = QDoubleSpinBox(); self.spalling.setRange(0,100); self.spalling.setDecimals(1); self.spalling.setValue(self.cave.spalling_pct); self._set_uniform_input_width(self.spalling); g.addWidget(QLabel("% spalling as fines"), row,0); g.addWidget(self.spalling,row,1); row+=1
        return w

    def _build_primary_tab(self):
        w = QWidget()
        layout = QHBoxLayout(w)

        controls = QVBoxLayout()
        controls.addWidget(QLabel("<b>Primary run</b>"))
        self.nblocks = QSpinBox(); self.nblocks.setRange(100,200000); self.nblocks.setValue(20000)
        controls.addWidget(QLabel("Blocks to generate"))
        controls.addWidget(self.nblocks)
        self.btn_run_primary = QPushButton("Run primary")
        self.btn_save_prm = QPushButton("Save .PRM…"); self.btn_save_prm.setEnabled(False)
        hl = QHBoxLayout(); hl.addWidget(self.btn_run_primary); hl.addWidget(self.btn_save_prm)
        controls.addLayout(hl)
        controls.addStretch(1)

        layout.addLayout(controls, 0)

        plot_layout = QVBoxLayout()
        self.plot_primary = PlotWidget()
        plot_layout.addWidget(self.plot_primary)
        plot_layout.setStretch(plot_layout.indexOf(self.plot_primary), 1)
        btn_row = QHBoxLayout()
        btn_row.addStretch(1)
        self.btn_save_primary_plot = QPushButton("Save chart…")
        self.btn_save_primary_plot.setEnabled(False)
        btn_row.addWidget(self.btn_save_primary_plot)
        plot_layout.addLayout(btn_row)
        layout.addLayout(plot_layout, 1)

        self._update_monte_carlo_combination_controls()

        return w

    def _build_secondary_tab(self):
        w = QWidget()
        layout = QHBoxLayout(w)

        controls = QVBoxLayout()
        controls.addWidget(QLabel("<b>Secondary run & Hang-up</b>"))
        self.draw_height = QDoubleSpinBox(); self.draw_height.setRange(1,2000); self.draw_height.setValue(self.secondary.draw_height)
        controls.addWidget(QLabel("Draw height (m)")); controls.addWidget(self.draw_height)
        self.max_caving = QDoubleSpinBox(); self.max_caving.setRange(1,5000); self.max_caving.setValue(self.secondary.max_caving_height)
        controls.addWidget(QLabel("Max caving height (m)")); controls.addWidget(self.max_caving)
        self.swell = QDoubleSpinBox(); self.swell.setRange(1.0,3.0); self.swell.setSingleStep(0.05); self.swell.setValue(self.secondary.swell_factor)
        controls.addWidget(QLabel("Swell factor")); controls.addWidget(self.swell)
        self.draw_width = QDoubleSpinBox(); self.draw_width.setRange(1,200); self.draw_width.setValue(self.secondary.active_draw_width)
        controls.addWidget(QLabel("Active draw width (m)")); controls.addWidget(self.draw_width)
        self.add_fines = QDoubleSpinBox(); self.add_fines.setRange(0,80); self.add_fines.setValue(self.secondary.added_fines_pct)
        controls.addWidget(QLabel("Added fines % (dilution)")); controls.addWidget(self.add_fines)
        self.rate = QDoubleSpinBox(); self.rate.setRange(0,100); self.rate.setValue(self.secondary.rate_cm_day)
        controls.addWidget(QLabel("Draw rate (cm/day)")); controls.addWidget(self.rate)
        self.upper_w = QDoubleSpinBox(); self.upper_w.setRange(1,50); self.upper_w.setValue(self.secondary.drawbell_upper_width)
        controls.addWidget(QLabel("Drawbell upper width (m)")); controls.addWidget(self.upper_w)
        self.lower_w = QDoubleSpinBox(); self.lower_w.setRange(1,50); self.lower_w.setValue(self.secondary.drawbell_lower_width)
        controls.addWidget(QLabel("Drawbell lower width (m)")); controls.addWidget(self.lower_w)
        self.hang_method = QComboBox(); self.hang_method.addItems(["Ore-pass rules (width)","Kear method (area)"])
        controls.addWidget(QLabel("Hang-up method")); controls.addWidget(self.hang_method)
        self.btn_run_secondary = QPushButton("Run secondary + hang-up")
        self.btn_save_sec = QPushButton("Save .SEC…"); self.btn_save_sec.setEnabled(False)
        hl = QHBoxLayout(); hl.addWidget(self.btn_run_secondary); hl.addWidget(self.btn_save_sec)
        controls.addLayout(hl)
        controls.addStretch(1)

        layout.addLayout(controls, 0)

        plot_layout = QVBoxLayout()
        self.plot_secondary = PlotWidget()
        plot_layout.addWidget(self.plot_secondary)
        plot_layout.setStretch(plot_layout.indexOf(self.plot_secondary), 1)
        btn_row = QHBoxLayout(); btn_row.addStretch(1)
        self.btn_save_secondary_plot = QPushButton("Save chart…")
        self.btn_save_secondary_plot.setEnabled(False)
        btn_row.addWidget(self.btn_save_secondary_plot)
        plot_layout.addLayout(btn_row)
        self.lbl_hang = QLabel("Hang-up: -")
        self.lbl_hang.setAlignment(Qt.AlignLeft | Qt.AlignVCenter)
        plot_layout.addWidget(self.lbl_hang)
        layout.addLayout(plot_layout, 1)

        return w

    def _build_monte_carlo_tab(self):
        w = QWidget()
        layout = QHBoxLayout(w)

        controls = QVBoxLayout()
        controls.addWidget(QLabel("<b>Monte Carlo simulation</b>"))
        self.mc_runs = QSpinBox(); self.mc_runs.setRange(1, 10000); self.mc_runs.setValue(10)
        controls.addWidget(QLabel("Number of runs")); controls.addWidget(self.mc_runs)
        self.mc_blocks = QSpinBox(); self.mc_blocks.setRange(100, 50000); self.mc_blocks.setValue(5000)
        controls.addWidget(QLabel("Blocks per run")); controls.addWidget(self.mc_blocks)
        self.mc_variation = QDoubleSpinBox(); self.mc_variation.setRange(0.0, 100.0); self.mc_variation.setSingleStep(1.0); self.mc_variation.setValue(15.0)
        self.mc_variation.setSuffix(" %")
        controls.addWidget(QLabel("Parameter variation (±%)")); controls.addWidget(self.mc_variation)
        controls.addWidget(QLabel("Joint set combinations"))
        self.combo_mc_combinations = QComboBox()
        controls.addWidget(self.combo_mc_combinations)
        self.btn_run_monte_carlo = QPushButton("Run Monte Carlo")
        controls.addWidget(self.btn_run_monte_carlo)
        controls.addStretch(1)

        layout.addLayout(controls, 0)

        plot_layout = QVBoxLayout()
        toggle_row = QHBoxLayout()
        toggle_row.addWidget(QLabel("Series visibility:"))
        self.chk_show_primary = QCheckBox("Primary")
        self.chk_show_primary.setChecked(True)
        self.chk_show_secondary = QCheckBox("Secondary")
        self.chk_show_secondary.setChecked(True)
        toggle_row.addWidget(self.chk_show_primary)
        toggle_row.addWidget(self.chk_show_secondary)
        toggle_row.addStretch(1)
        plot_layout.addLayout(toggle_row)

        self.plot_monte_carlo = PlotWidget()
        plot_layout.addWidget(self.plot_monte_carlo)
        plot_layout.setStretch(plot_layout.indexOf(self.plot_monte_carlo), 1)
        btn_row = QHBoxLayout(); btn_row.addStretch(1)
        self.btn_save_mc_report = QPushButton("Save report…")
        self.btn_save_mc_report.setEnabled(False)
        btn_row.addWidget(self.btn_save_mc_report)
        self.btn_save_mc_plot = QPushButton("Save chart…")
        self.btn_save_mc_plot.setEnabled(False)
        btn_row.addWidget(self.btn_save_mc_plot)
        plot_layout.addLayout(btn_row)
        self.lbl_mc_stats = QLabel("Monte Carlo: -")
        plot_layout.addWidget(self.lbl_mc_stats)
        layout.addLayout(plot_layout, 1)

        return w

    def _build_chart_settings_tab(self):
        w = QWidget()
        layout = QVBoxLayout(w)

        self._chart_title_widgets.clear()
        self._axis_format_combos.clear()

        title_group = QGroupBox("Chart titles")
        title_layout = QGridLayout(title_group)
        title_defs = [
            ("primary", "Primary chart title", "Primary fragmentation (cum. mass)"),
            ("secondary", "Secondary chart title", "Primary vs Secondary (cum. mass)"),
            ("monte_carlo", "Monte Carlo chart title", "Monte Carlo cumulative mass envelopes"),
        ]
        for row, (key, label, default) in enumerate(title_defs):
            title_layout.addWidget(QLabel(label), row, 0)
            edit = QLineEdit(default)
            self._chart_title_widgets[key] = edit
            title_layout.addWidget(edit, row, 1)
            edit.textChanged.connect(self._refresh_all_plots)
        layout.addWidget(title_group)

        axis_group = QGroupBox("Axis number formatting")
        axis_layout = QHBoxLayout(axis_group)
        axis_layout.addWidget(QLabel("Volume axis"))
        fmt_options = ["Auto", "0", "0.0", "0.00", "Scientific (1eX)"]
        combo_x = QComboBox()
        combo_x.addItems(fmt_options)
        self._axis_format_combos["x"] = combo_x
        axis_layout.addWidget(combo_x)
        axis_layout.addSpacing(12)
        axis_layout.addWidget(QLabel("Cumulative axis"))
        combo_y = QComboBox()
        combo_y.addItems(fmt_options)
        self._axis_format_combos["y"] = combo_y
        axis_layout.addWidget(combo_y)
        axis_layout.addStretch(1)
        combo_x.currentIndexChanged.connect(self._refresh_all_plots)
        combo_y.currentIndexChanged.connect(self._refresh_all_plots)
        layout.addWidget(axis_group)

        series_group = QGroupBox("Series styling")
        series_layout = QVBoxLayout(series_group)
        series_layout.addWidget(QLabel("Configure legend label, color, dash, and width for each data series."))
        self._series_style_container = QVBoxLayout()
        series_layout.addLayout(self._series_style_container)
        series_layout.addStretch(1)
        layout.addWidget(series_group)

        layout.addStretch(1)
        return w

    def _build_defaults_tab(self):
        w = QWidget(); g = QGridLayout(w)
        row=0; g.addWidget(QLabel("<b>Defaults</b>"), row,0,1,2); row+=1
        self.lhd_cutoff = QDoubleSpinBox(); self.lhd_cutoff.setRange(0.1,50); self.lhd_cutoff.setValue(self.defaults.LHD_cutoff_m3); self._set_uniform_input_width(self.lhd_cutoff); g.addWidget(QLabel("LHD bucket cutoff (m³)"), row,0); g.addWidget(self.lhd_cutoff,row,1); row+=1
        self.seed = QSpinBox(); self.seed.setRange(0,10**9); self.seed.setValue(self.defaults.seed or 1234); self._set_uniform_input_width(self.seed); g.addWidget(QLabel("Random seed"), row,0); g.addWidget(self.seed,row,1); row+=1
        self.arching_pct = QDoubleSpinBox(); self.arching_pct.setRange(0,1); self.arching_pct.setSingleStep(0.01); self.arching_pct.setValue(self.defaults.arching_pct); self._set_uniform_input_width(self.arching_pct); g.addWidget(QLabel("Arching split fraction (0–1)"), row,0); g.addWidget(self.arching_pct,row,1); row+=1
        self.arch_factor = QDoubleSpinBox(); self.arch_factor.setRange(1,100); self.arch_factor.setValue(self.defaults.arch_stress_conc); self._set_uniform_input_width(self.arch_factor); g.addWidget(QLabel("Arch stress factor (× cave pressure)"), row,0); g.addWidget(self.arch_factor,row,1); row+=1
        return w

    def _connect_signals(self):
        self.btn_add_joint_set.clicked.connect(self.on_add_joint_set)
        self.btn_run_primary.clicked.connect(self.on_run_primary)
        self.btn_save_prm.clicked.connect(self.on_save_prm)
        self.btn_run_secondary.clicked.connect(self.on_run_secondary)
        self.btn_save_sec.clicked.connect(self.on_save_sec)
        self.btn_save_primary_plot.clicked.connect(lambda: self.plot_primary.save_dialog(self, "primary_distribution.png"))
        self.btn_save_secondary_plot.clicked.connect(lambda: self.plot_secondary.save_dialog(self, "secondary_distribution.png"))
        self.btn_run_monte_carlo.clicked.connect(self.on_run_monte_carlo)
        self.btn_save_mc_plot.clicked.connect(lambda: self.plot_monte_carlo.save_dialog(self, "monte_carlo.png"))
        self.btn_save_mc_report.clicked.connect(self.on_save_monte_carlo_report)
        self.chk_show_primary.toggled.connect(self._refresh_monte_carlo_plot)
        self.chk_show_secondary.toggled.connect(self._refresh_monte_carlo_plot)
        self.action_save_settings.triggered.connect(self.on_save_settings)
        self.action_load_settings.triggered.connect(self.on_load_settings)
        self.sig.done_primary.connect(self.on_done_primary)
        self.sig.done_secondary.connect(self.on_done_secondary)
        self.sig.done_hangup.connect(self.on_done_hangup)
        self.sig.done_monte_carlo.connect(self.on_done_monte_carlo)

    def update_models_from_ui(self):
        self.rock = RockMass(
            rock_type=self.rock_type.text() or "Unknown",
            MRMR=self.mrmr.value(),
            IRS=self.irs.value(),
            mi=self.mi.value(),
            frac_freq=self.ff.value(),
            frac_condition=self.fc.value(),
            density=self.density.value(),
        )
        sets = []
        for idx, entry in enumerate(self.joint_widgets):
            name_widget = entry.get("name")
            name = name_widget.text().strip() if isinstance(name_widget, QLineEdit) else ""
            if not name:
                name = f"Set{idx + 1}"
            dip = entry.get("dip")
            dip_range = entry.get("dip_range")
            dip_dir = entry.get("dip_dir")
            dip_dir_range = entry.get("dip_dir_range")
            jc = entry.get("jc")
            s_min = entry.get("spacing_min")
            s_mean = entry.get("spacing_mean")
            s_max = entry.get("spacing_max")
            sets.append(JointSet(
                name=name,
                mean_dip=dip.value() if isinstance(dip, QDoubleSpinBox) else 45.0,
                dip_range=dip_range.value() if isinstance(dip_range, QDoubleSpinBox) else 10.0,
                mean_dip_dir=dip_dir.value() if isinstance(dip_dir, QDoubleSpinBox) else 0.0,
                dip_dir_range=dip_dir_range.value() if isinstance(dip_dir_range, QDoubleSpinBox) else 20.0,
                spacing=SpacingDist(
                    "trunc_exp",
                    s_min.value() if isinstance(s_min, QDoubleSpinBox) else 0.3,
                    s_mean.value() if isinstance(s_mean, QDoubleSpinBox) else 1.0,
                    s_max.value() if isinstance(s_max, QDoubleSpinBox) else 3.0,
                ),
                JC=jc.value() if isinstance(jc, QSpinBox) else 20,
            ))
        self.joint_sets = sets
        self._update_combination_controls()
        self.cave = CaveFace(
            dip=self.cave_dip.value(),
            dip_dir=self.cave_ddir.value(),
            stress_dip=self.st_dip.value(),
            stress_strike=self.st_strike.value(),
            stress_normal=self.st_norm.value(),
            allow_stress_fractures=self.allow_sf.isChecked(),
            spalling_pct=self.spalling.value(),
        )
        self.defaults = Defaults(
            LHD_cutoff_m3=self.lhd_cutoff.value(),
            seed=int(self.seed.value()),
            arching_pct=self.arching_pct.value(),
            arch_stress_conc=self.arch_factor.value()
        )
        self.secondary = SecondaryRun(
            draw_height=self.draw_height.value(),
            max_caving_height=self.max_caving.value(),
            swell_factor=self.swell.value(),
            active_draw_width=self.draw_width.value(),
            added_fines_pct=self.add_fines.value(),
            rate_cm_day=self.rate.value(),
            drawbell_upper_width=self.upper_w.value(),
            drawbell_lower_width=self.lower_w.value(),
        )

    def _current_mc_selection(self) -> Optional[Tuple[str, Optional[Tuple[int, int, int]]]]:
        combo = getattr(self, "combo_mc_combinations", None)
        if isinstance(combo, QComboBox):
            data = combo.currentData()
            if isinstance(data, tuple) and len(data) == 2:
                mode, indexes = data
                if indexes is not None:
                    indexes = tuple(indexes)
                return mode, indexes
        return None

    def _serialize_mc_selection(self) -> Optional[dict]:
        selection = self._current_mc_selection()
        if not selection:
            return None
        mode, indexes = selection
        return {"mode": mode, "indexes": list(indexes) if indexes is not None else None}

    def _apply_mc_selection(self, selection: Optional[dict]):
        combo = getattr(self, "combo_mc_combinations", None)
        if not isinstance(combo, QComboBox):
            return
        if not selection:
            combo.setCurrentIndex(0)
            return
        mode = selection.get("mode")
        indexes = selection.get("indexes")
        if isinstance(indexes, list):
            indexes = tuple(int(i) for i in indexes)
        target = (mode, indexes if indexes is not None else None)
        idx = combo.findData(target)
        if idx >= 0:
            combo.setCurrentIndex(idx)
        else:
            combo.setCurrentIndex(0)

    def _gather_settings_state(self) -> dict:
        self.update_models_from_ui()
        selected_combo = self._selected_joint_combination_indexes()
        state = {
            "rock": asdict(self.rock),
            "joint_sets": [asdict(js) for js in self.joint_sets],
            "cave": asdict(self.cave),
            "defaults": asdict(self.defaults),
            "secondary": asdict(self.secondary),
            "primary": {"nblocks": int(self.nblocks.value()) if hasattr(self, "nblocks") else None},
            "monte_carlo": {
                "runs": int(self.mc_runs.value()) if hasattr(self, "mc_runs") else None,
                "blocks_per_run": int(self.mc_blocks.value()) if hasattr(self, "mc_blocks") else None,
                "variation_pct": float(self.mc_variation.value()) if hasattr(self, "mc_variation") else None,
                "selection": self._serialize_mc_selection(),
            },
            "selected_joint_combination": list(selected_combo) if selected_combo else None,
            "chart_titles": {
                key: widget.text()
                for key, widget in self._chart_title_widgets.items()
                if isinstance(widget, QLineEdit)
            },
            "axis_formats": {
                axis: combo.currentText()
                for axis, combo in self._axis_format_combos.items()
                if isinstance(combo, QComboBox)
            },
            "series_styles": {
                label: self._collect_series_style(label)
                for label in list(self._series_style_widgets.keys())
            },
        }
        return state

    def _apply_series_style(self, label: str, style: dict):
        if not style:
            return
        self._ensure_series_style_controls([label])
        widgets = self._series_style_widgets.get(label)
        if not widgets:
            return
        name_widget = widgets.get("name")
        color_combo = widgets.get("color")
        dash_combo = widgets.get("dash")
        width_spin = widgets.get("width")
        if isinstance(name_widget, QLineEdit) and "label" in style:
            prev = name_widget.blockSignals(True)
            name_widget.setText(str(style.get("label", label)))
            name_widget.blockSignals(prev)
        if isinstance(color_combo, QComboBox) and "color" in style:
            idx = color_combo.findData(style.get("color"))
            if idx >= 0:
                prev = color_combo.blockSignals(True)
                color_combo.setCurrentIndex(idx)
                color_combo.blockSignals(prev)
        if isinstance(dash_combo, QComboBox) and "linestyle" in style:
            idx = dash_combo.findData(style.get("linestyle"))
            if idx >= 0:
                prev = dash_combo.blockSignals(True)
                dash_combo.setCurrentIndex(idx)
                dash_combo.blockSignals(prev)
        if isinstance(width_spin, QDoubleSpinBox) and "linewidth" in style:
            prev = width_spin.blockSignals(True)
            width_spin.setValue(float(style.get("linewidth", width_spin.value())))
            width_spin.blockSignals(prev)

    def _update_ui_from_models(self):
        self.rock_type.setText(self.rock.rock_type)
        self.mrmr.setValue(self.rock.MRMR)
        self.irs.setValue(self.rock.IRS)
        self.mi.setValue(self.rock.mi)
        self.ff.setValue(self.rock.frac_freq)
        self.fc.setValue(self.rock.frac_condition)
        self.density.setValue(self.rock.density)

        self._rebuild_joint_set_widgets(self.joint_sets)

        self.cave_dip.setValue(self.cave.dip)
        self.cave_ddir.setValue(self.cave.dip_dir)
        self.st_dip.setValue(self.cave.stress_dip)
        self.st_strike.setValue(self.cave.stress_strike)
        self.st_norm.setValue(self.cave.stress_normal)
        self.allow_sf.setChecked(self.cave.allow_stress_fractures)
        self.spalling.setValue(self.cave.spalling_pct)

        self.lhd_cutoff.setValue(self.defaults.LHD_cutoff_m3)
        self.seed.setValue(self.defaults.seed or 0)
        self.arching_pct.setValue(self.defaults.arching_pct)
        self.arch_factor.setValue(self.defaults.arch_stress_conc)

        self.draw_height.setValue(self.secondary.draw_height)
        self.max_caving.setValue(self.secondary.max_caving_height)
        self.swell.setValue(self.secondary.swell_factor)
        self.draw_width.setValue(self.secondary.active_draw_width)
        self.add_fines.setValue(self.secondary.added_fines_pct)
        self.rate.setValue(self.secondary.rate_cm_day)
        self.upper_w.setValue(self.secondary.drawbell_upper_width)
        self.lower_w.setValue(self.secondary.drawbell_lower_width)

        if hasattr(self, "nblocks"):
            self.nblocks.setValue(max(100, int(self.nblocks.value())))

    def _apply_settings_state(self, state: dict):
        rock_data = state.get("rock")
        if isinstance(rock_data, dict):
            self.rock = RockMass(**rock_data)

        joint_sets_data = state.get("joint_sets")
        if isinstance(joint_sets_data, list) and joint_sets_data:
            rebuilt_sets: List[JointSet] = []
            for idx, js_dict in enumerate(joint_sets_data, start=1):
                spacing_dict = js_dict.get("spacing") if isinstance(js_dict, dict) else {}
                if not isinstance(spacing_dict, dict):
                    spacing_dict = {}
                spacing = SpacingDist(
                    spacing_dict.get("type", "trunc_exp"),
                    spacing_dict.get("min", 0.3),
                    spacing_dict.get("mean", spacing_dict.get("min", 0.3)),
                    spacing_dict.get("max_or_90pct", spacing_dict.get("mean", 1.0)),
                    spacing_dict.get("max_obs"),
                )
                rebuilt_sets.append(
                    JointSet(
                        name=js_dict.get("name", f"Set{idx}"),
                        mean_dip=js_dict.get("mean_dip", 45.0),
                        dip_range=js_dict.get("dip_range", 10.0),
                        mean_dip_dir=js_dict.get("mean_dip_dir", 0.0),
                        dip_dir_range=js_dict.get("dip_dir_range", 20.0),
                        spacing=spacing,
                        JC=js_dict.get("JC", 20),
                    )
                )
            if rebuilt_sets:
                self.joint_sets = rebuilt_sets

        cave_data = state.get("cave")
        if isinstance(cave_data, dict):
            self.cave = CaveFace(**cave_data)

        defaults_data = state.get("defaults")
        if isinstance(defaults_data, dict):
            self.defaults = Defaults(**defaults_data)

        secondary_data = state.get("secondary")
        if isinstance(secondary_data, dict):
            self.secondary = SecondaryRun(**secondary_data)

        primary_data = state.get("primary")
        if isinstance(primary_data, dict) and hasattr(self, "nblocks"):
            try:
                self.nblocks.setValue(int(primary_data.get("nblocks", self.nblocks.value())))
            except (TypeError, ValueError):
                pass

        self._update_ui_from_models()

        selected_combo = state.get("selected_joint_combination")
        if isinstance(selected_combo, list) and len(selected_combo) == 3:
            combo_tuple = tuple(int(i) for i in selected_combo)
            idx = self.combo_joint_combination.findData(combo_tuple)
            if idx >= 0:
                self.combo_joint_combination.setCurrentIndex(idx)

        mc_state = state.get("monte_carlo", {})
        if isinstance(mc_state, dict):
            if "runs" in mc_state and hasattr(self, "mc_runs"):
                try:
                    self.mc_runs.setValue(int(mc_state.get("runs", self.mc_runs.value())))
                except (TypeError, ValueError):
                    pass
            if "blocks_per_run" in mc_state and hasattr(self, "mc_blocks"):
                try:
                    self.mc_blocks.setValue(int(mc_state.get("blocks_per_run", self.mc_blocks.value())))
                except (TypeError, ValueError):
                    pass
            if "variation_pct" in mc_state and hasattr(self, "mc_variation"):
                try:
                    self.mc_variation.setValue(float(mc_state.get("variation_pct", self.mc_variation.value())))
                except (TypeError, ValueError):
                    pass
            self._apply_mc_selection(mc_state.get("selection"))

        chart_titles = state.get("chart_titles")
        if isinstance(chart_titles, dict):
            for key, text in chart_titles.items():
                widget = self._chart_title_widgets.get(key)
                if isinstance(widget, QLineEdit):
                    prev = widget.blockSignals(True)
                    widget.setText(str(text))
                    widget.blockSignals(prev)

        axis_formats = state.get("axis_formats")
        if isinstance(axis_formats, dict):
            for axis, choice in axis_formats.items():
                combo = self._axis_format_combos.get(axis)
                if isinstance(combo, QComboBox):
                    idx = combo.findText(str(choice))
                    if idx >= 0:
                        prev = combo.blockSignals(True)
                        combo.setCurrentIndex(idx)
                        combo.blockSignals(prev)

        series_styles = state.get("series_styles")
        if isinstance(series_styles, dict):
            for label, style in series_styles.items():
                if isinstance(style, dict):
                    self._apply_series_style(label, style)

        self.update_models_from_ui()
        self._refresh_all_plots()

    def on_save_settings(self):
        try:
            state = self._gather_settings_state()
        except Exception as exc:
            QMessageBox.critical(self, "Cannot save", f"Failed to gather settings:\n{exc}")
            return

        path, _ = QFileDialog.getSaveFileName(
            self,
            "Save settings",
            "bcf_settings.json",
            "JSON files (*.json);;All files (*)",
        )
        if not path:
            return
        if not path.lower().endswith(".json"):
            path += ".json"
        try:
            with open(path, "w", encoding="utf-8") as fh:
                json.dump(state, fh, indent=2)
        except OSError as exc:
            QMessageBox.critical(self, "Save failed", f"Could not save settings:\n{exc}")
            return
        QMessageBox.information(self, "Settings saved", f"Settings written to:\n{path}")

    def on_load_settings(self):
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Load settings",
            "bcf_settings.json",
            "JSON files (*.json);;All files (*)",
        )
        if not path:
            return
        try:
            with open(path, "r", encoding="utf-8") as fh:
                state = json.load(fh)
        except (OSError, json.JSONDecodeError) as exc:
            QMessageBox.critical(self, "Load failed", f"Could not load settings:\n{exc}")
            return

        if not isinstance(state, dict):
            QMessageBox.critical(self, "Load failed", "Settings file format is not recognized.")
            return

        try:
            self._apply_settings_state(state)
        except Exception as exc:
            QMessageBox.critical(self, "Load failed", f"Could not apply settings:\n{exc}")
            return

        QMessageBox.information(self, "Settings loaded", f"Settings loaded from:\n{path}")

    def on_add_joint_set(self):
        js = JointSet(
            name=f"Set{len(self.joint_widgets) + 1}",
            mean_dip=45.0,
            dip_range=10.0,
            mean_dip_dir=0.0,
            dip_dir_range=20.0,
            spacing=SpacingDist("trunc_exp", 0.3, 1.0, 3.0),
            JC=20,
        )
        self._add_joint_set_widget(js)

    def on_run_primary(self):
        self.update_models_from_ui()
        n = int(self.nblocks.value())
        selected_sets = self._selected_joint_sets()
        if len(selected_sets) < 3:
            QMessageBox.warning(self, "Joint sets required", "Select three joint sets for the analysis.")
            return
        self._last_secondary_stats = None
        self.plot_secondary.fig.clear()
        self.plot_secondary.canvas.draw_idle()
        self.plot_secondary.has_data = False
        self.btn_save_secondary_plot.setEnabled(False)
        def work():
            blocks = generate_primary_blocks(n, self.rock, selected_sets, self.cave, self.defaults, seed=self.defaults.seed or 1234)
            primary_fines_ratio = self.cave.spalling_pct/100.0
            out_path = os.path.join(os.getcwd(), "primary_output.prm")
            write_prm(out_path, self.rock, self.cave, blocks, primary_fines_ratio)
            self.sig.done_primary.emit(blocks, primary_fines_ratio, out_path)
        threading.Thread(target=work, daemon=True).start()

    def on_done_primary(self, blocks, fines_ratio, path):
        self.primary_blocks = blocks
        self.primary_fines_ratio = fines_ratio
        stats = distributions_from_blocks(blocks)
        self._last_primary_stats = stats
        self._ensure_series_style_controls(["Primary"])
        self._refresh_primary_plot()
        self.btn_save_prm.setEnabled(True)
        self.btn_save_primary_plot.setEnabled(True)
        QMessageBox.information(self, "Primary run complete", f"Generated {len(blocks)} blocks.\nTemporary .PRM written to:\n{path}")

    def on_save_prm(self):
        if not self.primary_blocks:
            QMessageBox.warning(self, "No primary run", "Run primary first.")
            return
        path,_ = QFileDialog.getSaveFileName(self, "Save .PRM", "run.prm", "PRM files (*.prm)")
        if path:
            write_prm(path, self.rock, self.cave, self.primary_blocks, self.primary_fines_ratio)
            QMessageBox.information(self, "Saved", f"Wrote {path}")

    def on_run_secondary(self):
        if not self.primary_blocks:
            QMessageBox.warning(self, "No primary blocks", "Run primary first.")
            return
        self.update_models_from_ui()
        selected_sets = self._selected_joint_sets()
        if len(selected_sets) < 3:
            QMessageBox.warning(self, "Joint sets required", "Select three joint sets for the analysis.")
            return
        def work():
            mu = average_scatter_deg_from_jointsets(selected_sets)
            sec_blocks, sec_fines_ratio = run_secondary(self.primary_blocks, self.rock, self.secondary, self.defaults, mu_scatter_deg=mu, primary_fines_ratio=self.primary_fines_ratio)
            out_path = os.path.join(os.getcwd(), "secondary_output.sec")
            write_sec(out_path, self.rock, self.cave, sec_blocks, self.primary_fines_ratio)
            self.sig.done_secondary.emit(sec_blocks, sec_fines_ratio, out_path)
        threading.Thread(target=work, daemon=True).start()

    def on_done_secondary(self, sec_blocks, sec_fines_ratio, path):
        self.secondary_blocks = sec_blocks
        prim_stats = distributions_from_blocks(self.primary_blocks)
        class P:
            def __init__(self,b): self.V,self.Omega,self.joints_inside=b.V,b.Omega,b.joints_inside
        sec_stats = distributions_from_blocks([P(b) for b in sec_blocks])
        self._last_primary_stats = prim_stats
        self._last_secondary_stats = sec_stats
        self._ensure_series_style_controls(["Primary", "Secondary"])
        self._refresh_secondary_plot()
        self.btn_save_sec.setEnabled(True)
        self.btn_save_secondary_plot.setEnabled(True)

        if self.hang_method.currentText().startswith("Ore-pass"):
            stats = orepass_hangups(sec_blocks, bell_width=self.secondary.drawbell_lower_width, seed=self.defaults.seed or 1234)
            self.sig.done_hangup.emit(stats)
        else:
            area = self.secondary.drawbell_lower_width * self.secondary.drawbell_upper_width
            stats = kear_hangups(sec_blocks, bell_area=area, seed=self.defaults.seed or 1234)
            self.sig.done_hangup.emit(stats)

        QMessageBox.information(self, "Secondary complete", f"Generated {len(sec_blocks)} secondary blocks.\nTemporary .SEC written to:\n{path}")

    def on_save_sec(self):
        if not self.secondary_blocks:
            QMessageBox.warning(self, "No secondary run", "Run secondary first.")
            return
        path,_ = QFileDialog.getSaveFileName(self, "Save .SEC", "run.sec", "SEC files (*.sec)")
        if path:
            write_sec(path, self.rock, self.cave, self.secondary_blocks, self.primary_fines_ratio)
            QMessageBox.information(self, "Saved", f"Wrote {path}")

    def on_done_hangup(self, stats: dict):
        self.lbl_hang.setText(f"Hang-ups → High: {stats['n_high']}  Low: {stats['n_low']}  Total hang-up tons (proxy): {stats['total_hangup_tons']:.1f} t")

    def on_run_monte_carlo(self):
        self.update_models_from_ui()
        runs = int(self.mc_runs.value())
        nblocks = int(self.mc_blocks.value())
        variation = float(self.mc_variation.value())

        if len(self.joint_sets) < 3:
            QMessageBox.warning(self, "Joint sets required", "Define at least three joint sets before running Monte Carlo.")
            return

        selected_indexes = self._selected_joint_combination_indexes()
        choice = self.combo_mc_combinations.currentData() if hasattr(self, "combo_mc_combinations") else ("geology", None)
        if not isinstance(choice, tuple) or len(choice) != 2:
            choice = ("geology", None)

        available_combos = self._all_joint_combinations()
        combos_to_run: List[Tuple[int, int, int]]
        mode, data = choice
        if mode == "combo" and data is not None:
            combos_to_run = [tuple(data)]
        elif mode == "all":
            combos_to_run = available_combos
        else:
            if selected_indexes is None:
                QMessageBox.warning(self, "Joint set selection", "Select three joint sets on the Geology tab to use for Monte Carlo.")
                return
            combos_to_run = [selected_indexes]

        combos_to_run = [tuple(combo) for combo in combos_to_run if len(combo) == 3]
        combos_to_run = [combo for combo in combos_to_run if all(0 <= idx < len(self.joint_sets) for idx in combo)]
        combos_to_run = list(dict.fromkeys(combos_to_run))
        if not combos_to_run:
            QMessageBox.warning(self, "Joint set selection", "No valid joint-set combinations are available for Monte Carlo.")
            return

        names = [js.name for js in self.joint_sets]
        combination_specs = [
            {"indexes": combo, "label": self._combination_label(combo, names)}
            for combo in combos_to_run
        ]

        total_runs = runs * len(combination_specs)
        base_inputs = {
            "rock": asdict(self.rock),
            "joint_sets": [asdict(js) for js in self.joint_sets],
            "cave": asdict(self.cave),
            "defaults": asdict(self.defaults),
            "secondary": asdict(self.secondary),
            "selected_combination": list(selected_indexes) if selected_indexes else None,
            "joint_combinations": [
                {"label": spec["label"], "indexes": list(spec["indexes"])} for spec in combination_specs
            ],
        }

        self._last_monte_carlo_result = None
        self._last_monte_carlo_inputs = dict(base_inputs)
        self._last_monte_carlo_settings = {
            "runs": runs,
            "total_runs": total_runs,
            "blocks_per_run": nblocks,
            "variation_pct": variation,
            "combination_mode": self.combo_mc_combinations.currentText() if hasattr(self, "combo_mc_combinations") else "Geology selection",
            "combinations": [spec["label"] for spec in combination_specs],
        }
        self.btn_save_mc_plot.setEnabled(False)
        self.btn_save_mc_report.setEnabled(False)
        self.lbl_mc_stats.setText("Monte Carlo: running…")
        self._refresh_monte_carlo_plot()

        def work():
            primary_results: Dict[str, List[dict]] = defaultdict(list)
            secondary_results: Dict[str, List[dict]] = defaultdict(list)
            primary_avg_volumes: Dict[str, List[float]] = defaultdict(list)
            secondary_avg_volumes: Dict[str, List[float]] = defaultdict(list)
            combo_order: List[str] = []

            max_workers = max(1, min(total_runs, os.cpu_count() or 4))
            seeds = [random.randint(0, 10**9) for _ in range(total_runs)]
            tasks = []
            seed_idx = 0
            for spec in combination_specs:
                for _ in range(runs):
                    tasks.append((seeds[seed_idx], spec))
                    seed_idx += 1

            def consume_result(label: str, primary_stats: dict, secondary_stats: dict):
                if label not in combo_order:
                    combo_order.append(label)
                primary_results[label].append(primary_stats)
                primary_avg_volumes[label].append(primary_stats.get("avg_volume", 0.0))
                secondary_results[label].append(secondary_stats)
                secondary_avg_volumes[label].append(secondary_stats.get("avg_volume", 0.0))

            if max_workers == 1:
                for seed, spec in tasks:
                    label, primary_stats, sec_stats = _monte_carlo_worker(
                        seed, nblocks, variation, base_inputs, spec["indexes"], spec["label"]
                    )
                    consume_result(label, primary_stats, sec_stats)
            else:
                try:
                    with ProcessPoolExecutor(max_workers=max_workers) as executor:
                        futures = [
                            executor.submit(
                                _monte_carlo_worker,
                                seed,
                                nblocks,
                                variation,
                                base_inputs,
                                spec["indexes"],
                                spec["label"],
                            )
                            for seed, spec in tasks
                        ]
                        for fut in as_completed(futures):
                            label, primary_stats, sec_stats = fut.result()
                            consume_result(label, primary_stats, sec_stats)
                except Exception as exc:
                    try:
                        self.sig.log.emit(f"Process pool failed ({exc}); falling back to threads.")
                    except Exception:  # pragma: no cover - logging is best effort
                        pass
                    with ThreadPoolExecutor(max_workers=max_workers) as executor:
                        futures = [
                            executor.submit(
                                _monte_carlo_worker,
                                seed,
                                nblocks,
                                variation,
                                base_inputs,
                                spec["indexes"],
                                spec["label"],
                            )
                            for seed, spec in tasks
                        ]
                        for fut in as_completed(futures):
                            label, primary_stats, sec_stats = fut.result()
                            consume_result(label, primary_stats, sec_stats)

            if not combo_order:
                return

            def envelope(stats_list: List[dict]):
                xs_local = [lo for lo, _ in stats_list[0]["bins"]]
                cum_mass_runs = [stats["cum_mass"] for stats in stats_list]
                avg_line_local = [statistics.mean(vals) for vals in zip(*cum_mass_runs)]
                min_line_local = [min(vals) for vals in zip(*cum_mass_runs)]
                max_line_local = [max(vals) for vals in zip(*cum_mass_runs)]
                return xs_local, avg_line_local, min_line_local, max_line_local

            multi_combo = len(combo_order) > 1
            primary_series: List[Tuple[str, List[float]]] = []
            primary_xs: List[float] = []
            for label in combo_order:
                stats_list = primary_results.get(label)
                if not stats_list:
                    continue
                xs_local, avg_line, min_line, max_line = envelope(stats_list)
                primary_xs = xs_local
                compact = self._compact_combo_label(label)
                prefix = "P"
                primary_series.append((f"{prefix} avg – {compact}", avg_line))
                primary_series.append((f"{prefix} min – {compact}", min_line))
                primary_series.append((f"{prefix} max – {compact}", max_line))
                if not multi_combo:
                    break

            secondary_series: List[Tuple[str, List[float]]] = []
            secondary_xs: List[float] = []
            for label in combo_order:
                stats_list = secondary_results.get(label)
                if not stats_list:
                    continue
                xs_local, avg_line, min_line, max_line = envelope(stats_list)
                secondary_xs = xs_local
                compact = self._compact_combo_label(label)
                prefix = "S"
                secondary_series.append((f"{prefix} avg – {compact}", avg_line))
                secondary_series.append((f"{prefix} min – {compact}", min_line))
                secondary_series.append((f"{prefix} max – {compact}", max_line))
                if not multi_combo:
                    break

            def summarize(values: List[float]) -> Dict[str, float]:
                if not values:
                    return {"mean_avg_volume": 0.0, "min_avg_volume": 0.0, "max_avg_volume": 0.0}
                return {
                    "mean_avg_volume": statistics.mean(values),
                    "min_avg_volume": min(values),
                    "max_avg_volume": max(values),
                }

            all_primary_avgs = [val for vals in primary_avg_volumes.values() for val in vals]
            all_secondary_avgs = [val for vals in secondary_avg_volumes.values() for val in vals]

            info = {
                "runs": total_runs,
                "variation_pct": variation,
                "runs_per_combination": runs,
                "primary": summarize(all_primary_avgs),
                "secondary": summarize(all_secondary_avgs),
                "combinations": [],
            }

            for label in combo_order:
                info["combinations"].append(
                    {
                        "label": label,
                        "indexes": list(next((spec["indexes"] for spec in combination_specs if spec["label"] == label), ())),
                        "runs": len(primary_avg_volumes.get(label, [])),
                        "primary": summarize(primary_avg_volumes.get(label, [])),
                        "secondary": summarize(secondary_avg_volumes.get(label, [])),
                    }
                )

            self.sig.done_monte_carlo.emit(
                {
                    "primary": {
                        "xs": primary_xs,
                        "series": primary_series,
                    },
                    "secondary": (
                        {
                            "xs": secondary_xs,
                            "series": secondary_series,
                        }
                        if secondary_series
                        else None
                    ),
                    "info": info,
                }
            )

        threading.Thread(target=work, daemon=True).start()

    def on_done_monte_carlo(self, result: dict):
        self._last_monte_carlo_result = result
        labels = []
        for key in ("primary", "secondary"):
            series = (result.get(key) or {}).get("series") or []
            labels.extend([label for label, _ in series])
        self._ensure_series_style_controls(labels)
        self._refresh_monte_carlo_plot()
        info = result.get("info", {})
        p_info = info.get("primary", {})
        s_info = info.get("secondary", {})
        combos_info = info.get("combinations", [])
        combos_summary = ""
        if combos_info:
            summary_parts = [f"{entry.get('label', 'Combo')} ({entry.get('runs', 0)} runs)" for entry in combos_info]
            combos_summary = " | Combinations: " + "; ".join(summary_parts)
            runs_per_combo = info.get("runs_per_combination")
            if runs_per_combo and len(combos_info) > 1:
                combos_summary += f" | Runs/combo: {runs_per_combo}"

        summary_text = (
            "Monte Carlo runs: {runs} | Primary avg volume {p_mean:.2f} m³ (min {p_min:.2f}, max {p_max:.2f}) | "
            "Secondary avg volume {s_mean:.2f} m³ (min {s_min:.2f}, max {s_max:.2f}) | Variation ±{var:.1f}%".format(
                runs=info.get("runs", 0),
                p_mean=p_info.get("mean_avg_volume", 0.0),
                p_min=p_info.get("min_avg_volume", 0.0),
                p_max=p_info.get("max_avg_volume", 0.0),
                s_mean=s_info.get("mean_avg_volume", 0.0),
                s_min=s_info.get("min_avg_volume", 0.0),
                s_max=s_info.get("max_avg_volume", 0.0),
                var=info.get("variation_pct", 0.0),
            )
        )
        self.lbl_mc_stats.setText(summary_text + combos_summary)
        self.btn_save_mc_report.setEnabled(True)

    def _refresh_all_plots(self):
        self._refresh_primary_plot()
        self._refresh_secondary_plot()
        self._refresh_monte_carlo_plot()

    def _refresh_primary_plot(self):
        stats = self._last_primary_stats
        if not stats:
            return
        xs = [lo for lo, _ in stats.get("bins", [])]
        ys = stats.get("cum_mass") or []
        if not xs or not ys:
            return
        label = "Primary"
        self._ensure_series_style_controls([label])
        styles = {label: self._collect_series_style(label)}
        self.plot_primary.plot_lines(
            xs,
            [ys],
            labels=[label],
            title=self._chart_title("primary", "Primary fragmentation (cum. mass)"),
            xlabel="Block volume (m³)",
            ylabel="Cumulative mass (%)",
            logx=True,
            styles=styles,
            x_formatter=self._axis_formatter_from_choice(self._axis_format_choice("x")),
            y_formatter=self._axis_formatter_from_choice(self._axis_format_choice("y")),
        )

    def _refresh_secondary_plot(self):
        prim_stats = self._last_primary_stats
        sec_stats = self._last_secondary_stats
        if not prim_stats:
            return
        staged: List[Tuple[str, List[float], List[float]]] = []
        labels: List[str] = []
        xs_primary = [lo for lo, _ in prim_stats.get("bins", [])]
        ys_primary = prim_stats.get("cum_mass") or []
        if xs_primary and ys_primary:
            staged.append(("Primary", xs_primary, ys_primary))
            labels.append("Primary")
        if sec_stats:
            xs_secondary = [lo for lo, _ in sec_stats.get("bins", [])]
            ys_secondary = sec_stats.get("cum_mass") or []
            if xs_secondary and ys_secondary:
                staged.append(("Secondary", xs_secondary, ys_secondary))
                labels.append("Secondary")
        if not staged:
            return
        self._ensure_series_style_controls(labels)
        xs_list = [entry[1] for entry in staged]
        ys_list = [entry[2] for entry in staged]
        style_lookup = {label: self._collect_series_style(label) for label in labels}
        self.plot_secondary.plot_lines(
            xs_list,
            ys_list,
            labels=labels,
            title=self._chart_title("secondary", "Primary vs Secondary (cum. mass)"),
            xlabel="Block volume (m³)",
            ylabel="Cumulative mass (%)",
            logx=True,
            styles=style_lookup,
            x_formatter=self._axis_formatter_from_choice(self._axis_format_choice("x")),
            y_formatter=self._axis_formatter_from_choice(self._axis_format_choice("y")),
        )

    def _axis_format_choice(self, axis: str) -> str:
        combo = self._axis_format_combos.get(axis)
        if isinstance(combo, QComboBox):
            return combo.currentText() or "Auto"
        return "Auto"

    def _chart_title(self, key: str, fallback: str) -> str:
        widget = self._chart_title_widgets.get(key)
        if isinstance(widget, QLineEdit):
            text = widget.text().strip()
            if text:
                return text
        return fallback

    def _refresh_monte_carlo_plot(self):
        result = self._last_monte_carlo_result
        xs_list: List[List[float]] = []
        lines: List[List[float]] = []
        labels: List[str] = []
        styles: Dict[str, Dict[str, object]] = {}
        staged_series: List[Tuple[str, List[float], List[float]]] = []
        label_names: List[str] = []

        if result and self.chk_show_primary.isChecked():
            primary = result.get("primary") or {}
            xs = primary.get("xs")
            series = primary.get("series") or []
            if xs and series:
                for label, values in series:
                    staged_series.append((label, xs, values))
                    label_names.append(label)

        if result and self.chk_show_secondary.isChecked():
            secondary = result.get("secondary") or {}
            xs = secondary.get("xs")
            series = secondary.get("series") or []
            if xs and series:
                for label, values in series:
                    staged_series.append((label, xs, values))
                    label_names.append(label)

        if label_names:
            self._ensure_series_style_controls(label_names)

        for label, xs, values in staged_series:
            xs_list.append(xs)
            lines.append(values)
            labels.append(label)
            styles[label] = self._collect_series_style(label)

        if lines:
            title_text = self._chart_title("monte_carlo", "Monte Carlo cumulative mass envelopes")
            self.plot_monte_carlo.plot_lines(
                xs_list,
                lines,
                labels=labels,
                title=title_text,
                xlabel="Block volume (m³)",
                ylabel="Cumulative mass (%)",
                logx=True,
                styles=styles,
                x_formatter=self._axis_formatter_from_choice(self._axis_format_choice("x")),
                y_formatter=self._axis_formatter_from_choice(self._axis_format_choice("y")),
            )
            self.btn_save_mc_plot.setEnabled(True)
        else:
            self.plot_monte_carlo.fig.clear()
            self.plot_monte_carlo.canvas.draw_idle()
            self.plot_monte_carlo.has_data = False
            self.btn_save_mc_plot.setEnabled(False)

    def on_save_monte_carlo_report(self):
        if not self._last_monte_carlo_result or not self._last_monte_carlo_settings:
            QMessageBox.warning(self, "No Monte Carlo run", "Run Monte Carlo before saving a report.")
            return
        path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Monte Carlo report",
            "monte_carlo_report.docx",
            "Word document (*.docx);;All files (*)",
        )
        if not path:
            return

        if not path.lower().endswith(".docx"):
            path += ".docx"

        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            QMessageBox.critical(
                self,
                "Missing dependency",
                "Saving Word reports requires the 'python-docx' package.\n"
                "Install it with 'pip install python-docx' and try again.",
            )
            return

        info = self._last_monte_carlo_result.get("info", {})
        primary = self._last_monte_carlo_result.get("primary") or {}
        secondary = self._last_monte_carlo_result.get("secondary") or {}
        doc = Document()
        doc.add_heading("Monte Carlo Fragmentation Report", 0)
        subtitle = doc.add_paragraph(datetime.now().strftime("Generated on %Y-%m-%d at %H:%M"))
        subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.add_heading("Simulation summary", level=1)
        summary_table = doc.add_table(rows=3, cols=4)
        header_cells = summary_table.rows[0].cells
        header_cells[0].text = "Stage"
        header_cells[1].text = "Average volume (m³)"
        header_cells[2].text = "Minimum (m³)"
        header_cells[3].text = "Maximum (m³)"

        primary_row = summary_table.rows[1].cells
        primary_row[0].text = "Primary"
        primary_row[1].text = f"{info.get('primary', {}).get('mean_avg_volume', 0.0):.3f}"
        primary_row[2].text = f"{info.get('primary', {}).get('min_avg_volume', 0.0):.3f}"
        primary_row[3].text = f"{info.get('primary', {}).get('max_avg_volume', 0.0):.3f}"

        secondary_row = summary_table.rows[2].cells
        secondary_row[0].text = "Secondary"
        secondary_row[1].text = f"{info.get('secondary', {}).get('mean_avg_volume', 0.0):.3f}"
        secondary_row[2].text = f"{info.get('secondary', {}).get('min_avg_volume', 0.0):.3f}"
        secondary_row[3].text = f"{info.get('secondary', {}).get('max_avg_volume', 0.0):.3f}"

        doc.add_paragraph(
            "Runs per combination: {runs}  |  Total runs: {total}  |  Blocks per run: {blocks}  |  Variation: ±{var:.1f}%".format(
                runs=self._last_monte_carlo_settings.get("runs", 0),
                total=self._last_monte_carlo_settings.get("total_runs", self._last_monte_carlo_settings.get("runs", 0)),
                blocks=self._last_monte_carlo_settings.get("blocks_per_run", 0),
                var=self._last_monte_carlo_settings.get("variation_pct", 0.0),
            )
        )

        combos_info = info.get("combinations", [])
        if combos_info:
            doc.add_heading("Joint set combinations", level=1)
            combo_table = doc.add_table(rows=len(combos_info) + 1, cols=6)
            headers = combo_table.rows[0].cells
            headers[0].text = "Combination"
            headers[1].text = "Indexes"
            headers[2].text = "Runs"
            headers[3].text = "Primary avg (m³)"
            headers[4].text = "Secondary avg (m³)"
            headers[5].text = "Ranges (m³)"
            for row_idx, combo in enumerate(combos_info, start=1):
                cells = combo_table.rows[row_idx].cells
                cells[0].text = combo.get("label", "")
                indexes = combo.get("indexes") or []
                if indexes:
                    cells[1].text = ", ".join(str(int(i) + 1) for i in indexes)
                else:
                    cells[1].text = "-"
                cells[2].text = str(combo.get("runs", 0))
                primary_stats = combo.get("primary", {})
                secondary_stats = combo.get("secondary", {})
                cells[3].text = f"{primary_stats.get('mean_avg_volume', 0.0):.3f}"
                cells[4].text = f"{secondary_stats.get('mean_avg_volume', 0.0):.3f}"
                p_min = primary_stats.get("min_avg_volume", 0.0)
                p_max = primary_stats.get("max_avg_volume", 0.0)
                s_min = secondary_stats.get("min_avg_volume", 0.0)
                s_max = secondary_stats.get("max_avg_volume", 0.0)
                cells[5].text = (
                    f"Primary: {p_min:.3f}–{p_max:.3f}; Secondary: {s_min:.3f}–{s_max:.3f}"
                )

        doc.add_heading("Simulation settings", level=1)
        settings_table = doc.add_table(rows=len(self._last_monte_carlo_settings) + 1, cols=2)
        settings_table.rows[0].cells[0].text = "Setting"
        settings_table.rows[0].cells[1].text = "Value"
        for idx, (key, value) in enumerate(self._last_monte_carlo_settings.items(), start=1):
            settings_table.rows[idx].cells[0].text = key.replace("_", " ").title()
            settings_table.rows[idx].cells[1].text = str(value)

        doc.add_heading("Base input snapshot", level=1)

        def add_dict_table(title: str, payload: dict):
            if not payload:
                doc.add_paragraph(f"{title}: (not available)")
                return
            doc.add_paragraph(title, style="List Bullet")
            table = doc.add_table(rows=len(payload) + 1, cols=2)
            table.rows[0].cells[0].text = "Parameter"
            table.rows[0].cells[1].text = "Value"
            for idx, (key, value) in enumerate(payload.items(), start=1):
                table.rows[idx].cells[0].text = key.replace("_", " ")
                table.rows[idx].cells[1].text = json.dumps(value, ensure_ascii=False) if isinstance(value, (dict, list)) else str(value)

        base_inputs = self._last_monte_carlo_inputs or {}
        add_dict_table("Rock mass", base_inputs.get("rock", {}))
        add_dict_table("Cave face", base_inputs.get("cave", {}))
        add_dict_table("Defaults", base_inputs.get("defaults", {}))
        add_dict_table("Secondary fragmentation", base_inputs.get("secondary", {}))

        joint_sets = base_inputs.get("joint_sets") or []
        if joint_sets:
            doc.add_heading("Joint sets", level=2)
            for idx, js in enumerate(joint_sets, start=1):
                doc.add_paragraph(f"Joint set {idx}: {js.get('name', 'Set')}", style="List Number")
                js_table = doc.add_table(rows=len(js) + 1, cols=2)
                js_table.rows[0].cells[0].text = "Parameter"
                js_table.rows[0].cells[1].text = "Value"
                for row_idx, (key, value) in enumerate(js.items(), start=1):
                    js_table.rows[row_idx].cells[0].text = key.replace("_", " ")
                    if isinstance(value, dict):
                        js_table.rows[row_idx].cells[1].text = json.dumps(value, ensure_ascii=False)
                    else:
                        js_table.rows[row_idx].cells[1].text = str(value)

        def add_envelope_section(title: str, xs: List[float], series: List[tuple[str, List[float]]]):
            doc.add_heading(title, level=1)
            if not xs or not series:
                doc.add_paragraph("No data available.")
                return
            table = doc.add_table(rows=len(xs) + 1, cols=len(series) + 1)
            table.rows[0].cells[0].text = "Block volume (m³)"
            for idx, (label, _) in enumerate(series, start=1):
                table.rows[0].cells[idx].text = label
            for row_idx, vol in enumerate(xs, start=1):
                table.rows[row_idx].cells[0].text = f"{vol:.4f}"
                for col_idx, (_, values) in enumerate(series, start=1):
                    try:
                        table.rows[row_idx].cells[col_idx].text = f"{values[row_idx-1]:.3f}"
                    except IndexError:
                        table.rows[row_idx].cells[col_idx].text = "-"

        add_envelope_section(
            "Primary cumulative-mass envelope",
            primary.get("xs") or [],
            primary.get("series") or [],
        )
        add_envelope_section(
            "Secondary cumulative-mass envelope",
            secondary.get("xs") or [],
            secondary.get("series") or [],
        )

        try:
            doc.save(path)
        except Exception as exc:  # pragma: no cover - UI feedback
            QMessageBox.critical(self, "Save failed", f"Could not save report:\n{exc}")
        else:
            QMessageBox.information(self, "Report saved", f"Monte Carlo report saved to:\n{path}")

    def _ensure_series_style_controls(self, labels: List[str]):
        if not hasattr(self, "_series_style_container"):
            return
        for label in labels:
            if label not in self._series_style_widgets:
                self._add_series_style_row(label)

    def _add_series_style_row(self, label: str):
        row = QWidget()
        layout = QHBoxLayout(row)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(10)

        base_label = QLabel(label)
        base_label.setMinimumWidth(180)
        base_label.setStyleSheet("font-weight: 600; color: #1d2433;")
        layout.addWidget(base_label)

        name_edit = QLineEdit(label)
        name_edit.setPlaceholderText("Legend label")
        layout.addWidget(name_edit)

        color_combo = QComboBox()
        for color_name, color_value in MC_COLOR_OPTIONS:
            color_combo.addItem(color_name, color_value)
        default_color_index = self._preferred_color_index(label)
        layout.addWidget(color_combo)

        dash_combo = QComboBox()
        for text, pattern in LINE_STYLE_OPTIONS:
            dash_combo.addItem(text, pattern)
        layout.addWidget(dash_combo)

        width_spin = QDoubleSpinBox()
        width_spin.setRange(0.5, 8.0)
        width_spin.setSingleStep(0.1)
        layout.addWidget(width_spin)

        if self._should_randomize_series(label):
            color_idx, dash_idx, width_val = self._randomized_series_defaults(label)
        else:
            color_idx = default_color_index
            dash_idx = self._index_for_dash(self._default_dash_for_label(label))
            width_val = self._default_width_for_label(label)

        color_combo.setCurrentIndex(color_idx)
        dash_combo.setCurrentIndex(dash_idx)
        width_spin.setValue(width_val)

        layout.addStretch(1)
        self._series_style_container.addWidget(row)
        self._series_style_widgets[label] = {
            "container": row,
            "name": name_edit,
            "color": color_combo,
            "dash": dash_combo,
            "width": width_spin,
        }

        name_edit.textChanged.connect(self._refresh_all_plots)
        color_combo.currentIndexChanged.connect(self._refresh_all_plots)
        dash_combo.currentIndexChanged.connect(self._refresh_all_plots)
        width_spin.valueChanged.connect(self._refresh_all_plots)

    def _collect_series_style(self, label: str) -> Dict[str, object]:
        widgets = self._series_style_widgets.get(label)
        if not widgets:
            return {"label": label}
        name_widget = widgets.get("name")
        color_combo = widgets.get("color")
        dash_combo = widgets.get("dash")
        width_spin = widgets.get("width")
        custom_label = name_widget.text().strip() if isinstance(name_widget, QLineEdit) else label
        if not custom_label:
            custom_label = label
        color = color_combo.currentData() if isinstance(color_combo, QComboBox) else None
        linestyle = dash_combo.currentData() if isinstance(dash_combo, QComboBox) else "-"
        linewidth = width_spin.value() if isinstance(width_spin, QDoubleSpinBox) else 1.5
        return {
            "label": custom_label,
            "color": color,
            "linestyle": linestyle or "-",
            "linewidth": linewidth,
        }

    def _should_randomize_series(self, label: str) -> bool:
        text = label.lower()
        return " – " in label and (text.startswith("p ") or text.startswith("s "))

    def _series_combo_key(self, label: str) -> str:
        if "–" in label:
            _, _, tail = label.partition("–")
            return tail.strip()
        return label

    def _combo_role_from_label(self, label: str) -> str:
        text = label.lower()
        if "avg" in text:
            return "avg"
        if " min" in text or "min " in text:
            return "min"
        if " max" in text or "max " in text:
            return "max"
        return "line"

    def _randomized_series_defaults(self, label: str) -> Tuple[int, int, float]:
        cached = self._series_random_defaults.get(label)
        if cached:
            return cached
        combo_key = self._series_combo_key(label)
        cache = self._combo_style_cache.get(combo_key)
        if not cache:
            seed = abs(hash(combo_key)) & 0xFFFFFFFF
            rng = random.Random(seed)
            color_choices = list(range(len(MC_COLOR_OPTIONS)))
            if len(color_choices) > 2:
                color_choices = color_choices[2:]
            color_idx = color_choices[rng.randrange(len(color_choices))] if color_choices else 0
            dash_palettes = [
                {"avg": "-", "min": "--", "max": "-."},
                {"avg": "-", "min": ":", "max": "--"},
                {"avg": "-", "min": "--", "max": ":"},
                {"avg": "-", "min": "-.", "max": ":"},
            ]
            palette = dash_palettes[rng.randrange(len(dash_palettes))]
            cache = {"color": color_idx, "palette": palette, "seed": seed}
            self._combo_style_cache[combo_key] = cache
        role = self._combo_role_from_label(label)
        palette = cache.get("palette", {})
        dash = palette.get(role, "-")
        role_seed = cache.get("seed", 0) + (hash(role) & 0xFFFF)
        role_rng = random.Random(role_seed)
        if role == "avg":
            width = round(2.0 + role_rng.random() * 0.6, 1)
        elif role == "min":
            width = round(1.0 + role_rng.random() * 0.4, 1)
        elif role == "max":
            width = round(1.1 + role_rng.random() * 0.5, 1)
        else:
            width = round(1.3 + role_rng.random() * 0.4, 1)
        dash_idx = self._index_for_dash(dash)
        defaults = (cache.get("color", 0), dash_idx, width)
        self._series_random_defaults[label] = defaults
        return defaults

    def _default_dash_for_label(self, label: str) -> str:
        text = label.lower()
        if "minimum" in text or " min" in text or text.endswith(" min"):
            return "--"
        if "maximum" in text or " max" in text or text.endswith(" max"):
            return "-."
        return "-"

    def _default_width_for_label(self, label: str) -> float:
        text = label.lower()
        if "average" in text or "avg" in text:
            return 2.0
        if any(word in text for word in ("minimum", "maximum", " min", " max")):
            return 1.3
        return 1.5

    def _preferred_color_index(self, label: str) -> int:
        if self._should_randomize_series(label):
            return 0
        text = label.lower()
        if "primary" in text:
            if self._series_color_cursor < 2:
                self._series_color_cursor = 2
            return 0
        if "secondary" in text:
            if self._series_color_cursor < 2:
                self._series_color_cursor = 2
            return 1
        idx = self._series_color_cursor % len(MC_COLOR_OPTIONS)
        self._series_color_cursor += 1
        return idx

    def _index_for_dash(self, dash: str) -> int:
        for idx, (_, pattern) in enumerate(LINE_STYLE_OPTIONS):
            if pattern == dash:
                return idx
        return 0

    def _axis_formatter_from_choice(self, choice: str):
        option = (choice or "Auto").strip()
        if option == "Auto":
            return None
        if option == "0":
            return StrMethodFormatter("{x:.0f}")
        if option == "0.0":
            return StrMethodFormatter("{x:.1f}")
        if option == "0.00":
            return StrMethodFormatter("{x:.2f}")
        if option.startswith("Scientific"):
            formatter = ScalarFormatter(useOffset=False, useMathText=True)
            formatter.set_scientific(True)
            formatter.set_powerlimits((-3, 4))
            return formatter
        return None

def launch():
    app = QApplication([])
    mw = MainWindow()
    mw.show()
    app.exec()
